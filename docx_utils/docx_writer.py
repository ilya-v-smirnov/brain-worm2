from __future__ import annotations

import json
import re
from typing import List, Tuple
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from pathlib import Path

from dataclasses import dataclass

@dataclass(frozen=True)
class DocxStyleProfile:
    # Title (H1)
    h1_space_before: int = 0
    h1_space_after: int = 2
    h1_line_spacing: float = 1.0

    # Section headers
    h2_space_before: int = 6
    h2_space_after: int = 2
    h3_space_before: int = 4
    h3_space_after: int = 0

    # Body paragraphs
    body_align: int = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_space_after: int = 6         # больше воздуха между абзацами
    body_line_spacing: float = 1.0

    # Meta lines (e.g., Source path)
    meta_align: int = WD_ALIGN_PARAGRAPH.LEFT
    meta_space_after: int = 2

    # Figures
    figure_align: int = WD_ALIGN_PARAGRAPH.JUSTIFY
    figure_space_after: int = 6

DEFAULT_STYLE = DocxStyleProfile()


def _apply_paragraph_style(p, *, align=None, space_before=None, space_after=None, line_spacing=None) -> None:
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _apply_h1(p, style: DocxStyleProfile) -> None:
    _apply_paragraph_style(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=style.h1_space_before,
        space_after=style.h1_space_after,
        line_spacing=style.h1_line_spacing,
    )


def _apply_meta(p, style: DocxStyleProfile) -> None:
    _apply_paragraph_style(
        p,
        align=style.meta_align,
        space_before=0,
        space_after=style.meta_space_after,
        line_spacing=1.0,
    )


def _apply_body(p, style: DocxStyleProfile) -> None:
    _apply_paragraph_style(
        p,
        align=style.body_align,
        space_before=0,
        space_after=style.body_space_after,
        line_spacing=style.body_line_spacing,
    )


def _apply_figure(p, style: DocxStyleProfile) -> None:
    _apply_paragraph_style(
        p,
        align=style.figure_align,
        space_before=0,
        space_after=style.figure_space_after,
        line_spacing=1.0,
    )


_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_runs_with_bold_markdown(p, text: str) -> None:
    """
    Adds runs to an existing paragraph, converting **bold** to bold runs.
    Minimal parser: no nesting, no escaping.
    """
    s = str(text or "")
    pos = 0
    for m in _MD_BOLD_RE.finditer(s):
        if m.start() > pos:
            p.add_run(s[pos:m.start()])
        r = p.add_run(m.group(1))
        r.bold = True
        pos = m.end()
    if pos < len(s):
        p.add_run(s[pos:])


def _add_markdown_body(doc: Document, text: str, style: DocxStyleProfile) -> None:
    """
    Renders body text with minimal markdown support:
      - **bold** -> bold runs
      - lines starting with '* ' -> bullet points
    Paragraph boundaries are taken from _normalize_word_breaks() -> split by blank lines.
    """
    body = _normalize_word_breaks(str(text or "")).strip()
    if not body:
        body = "—"

    for chunk in body.split("\n\n"):
        chunk = chunk.strip()
        if not chunk:
            continue

        # A chunk may still contain multiple lines (rare). Handle line-by-line.
        lines = chunk.splitlines() or [chunk]
        for ln in lines:
            ln = ln.rstrip()

            if ln.lstrip().startswith("* "):
                bullet_text = ln.lstrip()[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                _apply_body(p, style)
                _add_runs_with_bold_markdown(p, bullet_text)
            else:
                p = doc.add_paragraph()
                _apply_body(p, style)
                _add_runs_with_bold_markdown(p, ln.strip())


# --- Базовые помощники -------------------------------------------------------

def _normalize_word_breaks(text: str) -> str:
    """
    Нормализует Word-переносы в "абзацы" (paragraph breaks).

    Что ловим:
    - literal token "^l" -> абзац
    - VT / vertical tab '\x0b' -> абзац
    - Unicode line/paragraph separators: U+2028, U+2029 -> абзац
    - CRLF/CR -> LF
    - одиночные '\n' внутри "сплошного" текста часто означают ^l: конвертируем в абзацы,
      но сохраняем уже существующие двойные пустые строки как разделители абзацев.
    """
    if not text:
        return ""
    s = str(text)

    # Normalize newlines
    s = s.replace("\r\n", "\n").replace("\r", "\n")

    # Word-like manual breaks
    s = s.replace("^l", "\n\n")
    s = s.replace("\x0b", "\n\n")         # vertical tab
    s = s.replace("\u2028", "\n\n")       # line separator
    s = s.replace("\u2029", "\n\n")       # paragraph separator

    # Если текст содержит одиночные переносы строк, делаем их абзацами.
    # При этом не ломаем уже имеющиеся двойные переносы.
    # Принцип:
    #   - сначала временно защищаем \n\n
    #   - затем все одиночные \n превращаем в \n\n
    #   - возвращаем защищённые
    sentinel = "\uFFFF"  # unlikely character
    s = s.replace("\n\n", sentinel)
    s = s.replace("\n", "\n\n")
    s = s.replace(sentinel, "\n\n")

    return s


def _p(doc: Document, text: str = "", *, bold: bool = False, size: int = 11, font: str | None = None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bool(bold)
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
    # === ВЫРАВНИВАНИЕ ПО ШИРИНЕ ===
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    return run

def _blank(doc: Document, n: int = 1):
    # максимум одна визуальная пустая строка
    n = 1 if n and n > 0 else 0
    for _ in range(n):
        p = doc.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)

def _heading(doc: Document, text: str):
    # Совместимость со старыми вызовами: заголовок уровня 2
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_after = Pt(0)

def _heading_h2(doc: Document, text: str):
    # настоящий заголовок Word уровня 2
    h = doc.add_heading(text, level=2)
    # уберём лишние интервалы после
    h.paragraph_format.space_after = Pt(0)

def _bullet_list(doc: Document, items):
    for it in items or []:
        p = doc.add_paragraph(str(it), style="List Bullet")
        # === ВЫРАВНИВАНИЕ ПО ШИРИНЕ ДЛЯ СПИСКОВ ===
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)

def _sections_block(doc: Document, sections: List[Dict[str, str]], style: DocxStyleProfile = DEFAULT_STYLE):
    """
    Рендерит список секций.
    Элемент: {"title": str, "text": Optional[str], "level": int, "suppress_empty_dash": bool=False}
      - text=None  -> вывести только заголовок (без абзаца)
      - text==""   -> вывести "—" (кроме suppress_empty_dash=True)
    """
    if not sections:
        return
    for sec in sections:
        title = (sec.get("title") or "Section").strip()
        text = sec.get("text", "")
        level = int(sec.get("level") or 2)
        suppress_empty = bool(sec.get("suppress_empty_dash", False))

        # Заголовок
        h = doc.add_heading(title, level=level)
        if level == 3:
            for r in h.runs:
                r.font.size = Pt(12)
        h.paragraph_format.space_after = Pt(0)

        # Тело
        printed_body = False
        if text is None:
            # только заголовок (например, Results, когда есть подпункты) — тела нет
            printed_body = False
        else:
            body = text.strip()
            if not body and suppress_empty:
                # ничего не печатаем
                printed_body = False
            else:
                if not body:
                    body = "—"
                
                body = _normalize_word_breaks(body)

                for chunk in body.split("\n\n"):
                    p = doc.add_paragraph(chunk.strip())
                    _apply_body(p, style)
                printed_body = True

        # Пустую строку добавляем только если реально печатали тело
        if printed_body:
            _blank(doc, 1)

def _abbrev_simple_table(doc: Document, pairs: List[tuple[str, str]]):
    """
    Рисует ТОЛЬКО таблицу (без заголовка). Интервалы после строк = 0 pt.
    """
    if not pairs:
        return
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Abbreviation"
    hdr[1].text = "Expanded"
    # Убираем интервалы после абзацев в ячейках заголовка
    for cell in hdr:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)

    for abbr, expanded in pairs:
        row = table.add_row().cells
        row[0].text = abbr or ""
        row[1].text = expanded or ""
        # Убираем интервалы после абзацев в каждой ячейке
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)

def _parse_abbreviation_pairs(text: str) -> List[Tuple[str, str]]:
    """
    Parses Variant A:
      ABBR — definition
    Accepts separators: em dash (—), en dash (–), hyphen (-), colon (:).
    Ignores empty and non-matching lines.
    """
    pairs: List[Tuple[str, str]] = []
    if not text:
        return pairs

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue

        m = re.match(r"^\s*(.+?)\s*(?:—|–|-|:)\s*(.+?)\s*$", line)
        if not m:
            continue

        abbr = m.group(1).strip()
        definition = m.group(2).strip()
        if abbr and definition:
            pairs.append((abbr, definition))

    return pairs

def _set_table_borders_none(table) -> None:
    """
    Removes visible borders for a python-docx table by setting tblBorders to 'nil'.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tblBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tblBorders.append(element)
        element.set(qn("w:val"), "nil")


def _abbrev_simple_table(doc: Document, pairs: List[Tuple[str, str]]):
    """
    Draws ONLY the table (no H2 header here).
    Borders are removed (transparent).
    """
    if not pairs:
        return

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    _set_table_borders_none(table)

    hdr = table.rows[0].cells
    hdr[0].text = "Abbreviation"
    hdr[1].text = "Definition"
    for cell in hdr:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)
            # make header bold
            for run in para.runs:
                run.bold = True

    for abbr, expanded in pairs:
        row = table.add_row().cells
        row[0].text = abbr or ""
        row[1].text = expanded or ""
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)


def _add_page_break_if_needed(doc: Document) -> None:
    if len(doc.paragraphs) > 0 or len(doc.tables) > 0:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def _write_figure_summaries(doc: Document, figure_summaries):
    """
    Печатает раздел 'Figure summaries' перед 'Abbreviations'.
    Ожидает список объектов вида: [{"figure": "Figure 1", "summary": "..."}]
    """
    if not figure_summaries:
        return
    _heading_h2(doc, "Figure summaries")
    for item in figure_summaries:
        fig = str(item.get("figure") or "").strip()
        summ = _normalize_word_breaks(str(item.get("summary") or "")).strip()
        if not fig or not summ:
            continue
        p = doc.add_paragraph()
        r = p.add_run(f"{fig}. ")
        r.bold = True
        p.add_run(summ)


def _loc(label_en: str, lang: str) -> str:
    """
    Простая локализация заголовков для RU.
    Мы трогаем только то, что требуется задачей: Results -> Результаты.
    Остальные заголовки оставляем как есть (минимальная правка).
    """
    lang = (lang or "").upper()
    if lang == "RU":
        mapping = {
            "Results": "Результаты",
        }
        return mapping.get(label_en, label_en)
    return label_en


def append_ai_summary_to_docx(
    *,
    docx_path: Path,
    summary: dict,
    style: DocxStyleProfile = DEFAULT_STYLE,
):
    """
    Добавляет одну версию AI-summary в docx.
    Если файл существует — page break.
    """

    if docx_path.exists():
        doc = Document(str(docx_path))
        _add_page_break_if_needed(doc)
    else:
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()

    header = summary.get("header") or {}

    # === HEADER ===
    h1 = doc.add_heading(f'{header.get("year","")} {header.get("title","")}', level=1)
    _apply_h1(h1, style)

    def meta_line(label, value):
        p = doc.add_paragraph()
        _apply_meta(p, style)
        r1 = p.add_run(f"{label}: "); r1.bold = True
        p.add_run(_normalize_word_breaks(str(value or "")))

    meta_line("Source path", header.get("source_path",""))
    meta_line("Model", header.get("model",""))
    meta_line("Language", header.get("language",""))
    _blank(doc, 1)

    # === KEY POINTS ===
    _heading_h2(doc, "Key points")
    _bullet_list(doc, summary.get("key_points"))
    _blank(doc, 1)

    # === SECTIONS ===
    sections_out = []

    sections_out.append({
        "title": "Introduction",
        "text": summary.get("introduction") or "—",
        "level": 2,
    })

    # Results: строго по подразделам
    results = summary.get("results") or []
    if results:
        sections_out.append({
            "title": "Results",
            "text": None,
            "level": 2,
            "suppress_empty_dash": True,
        })
        for r in results:
            sections_out.append({
                "title": r["section_title"],   # оригинальное название
                "text": r["mini_summary"],
                "level": 3,
            })

    sections_out.append({
        "title": "Discussion",
        "text": summary.get("discussion") or "—",
        "level": 2,
    })

    _sections_block(doc, sections_out, style=style)

    # === FIGURES ===
    _write_figure_summaries(doc, summary.get("figures", {}).get("items"))

    # === ABBREVIATIONS (если появятся позже) ===
    abbr = summary.get("abbreviations") or []
    if abbr:
        _heading_h2(doc, "Abbreviations")
        pairs = [(a["abbr"], a["expanded"]) for a in abbr]
        _abbrev_simple_table(doc, pairs)

    doc.save(str(docx_path))

    # --- Debug artifact: save the exact summary JSON next to the docx ---
    # Example: my_summary.docx -> my_summary.summary.json
    try:
        json_path = docx_path.with_suffix(".summary.json")
        json_path.write_text(
            json.dumps(summary, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        # Do not fail docx saving if JSON dump fails
        pass


# =============================================================================
# Export extracted article text (NOT AI-summary)
# =============================================================================

def export_extracted_text_to_docx(
    *,
    docx_path,
    article: dict,
    source_path: str = "",
    style: DocxStyleProfile = DEFAULT_STYLE,
):
    """
    Exports extracted article text JSON into a .docx file, using the same
    formatting conventions as this module.

    Header: standard "year title" + Source path (NO Model / Language).
    Sections order:
      Introduction -> Methods -> Results (subsections in JSON order) ->
      Discussion -> Figures (captions)
    """

    # Create new doc (overwrite / create)
    docx_path = Path(str(docx_path))
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    title = str(article.get("title") or "").strip()
    year = str(article.get("year") or "").strip()

    # === HEADER ===
    h1 = doc.add_heading(f"{year} {title}".strip(), level=1)
    _apply_h1(h1, style)
    
    def meta_line(label, value):
        p = doc.add_paragraph()
        _apply_meta(p, style)
        r1 = p.add_run(f"{label}: "); r1.bold = True
        p.add_run(_normalize_word_breaks(str(value or "")))

    if source_path:
        meta_line("Source path", source_path)

    _blank(doc, 1)

    # === SECTIONS ===
    def add_section_h2(name: str):
        _heading_h2(doc, name)

    def add_body(text: str):
        body = _normalize_word_breaks(str(text or "")).strip()
        if not body:
            body = "—"
        for chunk in body.split("\n\n"):
            p = doc.add_paragraph(chunk.strip())
            _apply_body(p, style)
        _blank(doc, 1)


    add_section_h2("Introduction")
    add_body(str(article.get("introduction") or ""))

    add_section_h2("Methods")
    add_body(str(article.get("methods") or ""))

    # Results: keep JSON order
    results = article.get("results") or []
    add_section_h2(_loc("Results", "EN"))
    if isinstance(results, list) and results:
        for item in results:
            if not isinstance(item, dict):
                continue
            sec_title = str(item.get("section_title") or "").strip()
            sec_text = str(item.get("section_text") or "").strip()

            if sec_title:
                h = doc.add_heading(sec_title, level=3)
                for r in h.runs:
                    r.font.size = Pt(12)
                h.paragraph_format.space_after = Pt(0)

            if sec_text:
                sec_text = _normalize_word_breaks(sec_text)
                if sec_text:
                    sec_text = _normalize_word_breaks(sec_text).strip()
                    for chunk in sec_text.split("\n\n"):
                        p = doc.add_paragraph(chunk.strip())
                        _apply_body(p, style)
                    _blank(doc, 1)


            else:
                # If there is a title but no text, print dash (consistent with module)
                if sec_title:
                    p = doc.add_paragraph("—")
                    _apply_body(p, style)
                    _blank(doc, 1)

    else:
        add_body("")

    add_section_h2("Discussion")
    add_body(str(article.get("discussion") or ""))

    # Figures
    figures = article.get("figures") or []
    add_section_h2("Figures")
    if isinstance(figures, list) and figures:
        for fig in figures:
            if not isinstance(fig, dict):
                continue
            num = str(fig.get("figure_number") or "").strip()
            cap = str(fig.get("figure_caption") or "").strip()
            if not (num or cap):
                continue

            p = doc.add_paragraph("—")
            _apply_body(p, style)

            _apply_figure(p, style)

            if num:
                r = p.add_run(f"{num}. ")
                r.bold = True
            if cap:
                p.add_run(cap)
            _blank(doc, 1)
    else:
        add_body("")

    doc.save(str(docx_path))


def _parse_bullets(text: str) -> list[str]:
    """
    Превращает текст из textbox в список буллетов.
    Поддерживает строки вида:
      - item
      • item
      * item
      1) item
      1. item
    Если маркеров нет — режет по непустым строкам.
    """
    if not text:
        return []
    lines = [ln.strip() for ln in str(text).splitlines()]
    lines = [ln for ln in lines if ln]

    out: list[str] = []
    for ln in lines:
        ln = ln.strip()
        ln = re.sub(r"^[-•*]\s+", "", ln)
        ln = re.sub(r"^\d+\s*[\.\)]\s+", "", ln)
        out.append(ln.strip())

    # Если всё превратилось в пустоту — вернём пустой
    out = [x for x in out if x]
    return out


def append_semi_manual_summary_to_docx(
    *,
    docx_path: Path,
    payload: dict,
    style: DocxStyleProfile = DEFAULT_STYLE,
    overwrite: bool = False,
) -> None:
    """
    Сохраняет semi-manual summary в .docx.
    - overwrite=True: перезаписать файл
    - overwrite=False: если файл существует -> добавляет page break и дописывает новую версию
    Ожидаемый payload (минимум):
      {
        "title": str,
        "year": str|int,
        "language": "EN"|"RU"|...,
        "summary": {
            "key_points": str,
            "introduction": str,
            "methods": str,
            "results": [{"section_title": str, "summary_text": str}, ...],
            "discussion": str,
            "figure_narrative": str,
        },
        "prompts_used": {... optional ...},
        "source_json": str,
        "source_pdf": str,
      }
    """
    docx_path = Path(str(docx_path))
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    if overwrite and docx_path.exists():
        try:
            docx_path.unlink()
        except Exception:
            # если не смогли удалить — попробуем всё равно открыть как Document и перезаписать
            pass

    if docx_path.exists():
        doc = Document(str(docx_path))
        _add_page_break_if_needed(doc)
    else:
        doc = Document()

    title = str(payload.get("title") or "").strip()
    year = str(payload.get("year") or "").strip()
    lang = str(payload.get("language") or "").strip()

    # === HEADER ===
    h1 = doc.add_heading(f"{year} {title}".strip(), level=1)
    _apply_h1(h1, style)

    def meta_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        _apply_meta(p, style)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        p.add_run(_normalize_word_breaks(str(value or "")))

    if payload.get("source_pdf"):
        meta_line("Source", str(payload.get("source_pdf") or ""))
    if lang:
        meta_line("Language", lang)

    _blank(doc, 1)

    summary = payload.get("summary") or {}

    # === KEY POINTS ===
    _heading_h2(doc, "Key points")
    kp_items = _parse_bullets(str(summary.get("key_points") or ""))
    if kp_items:
        _bullet_list(doc, kp_items)
    else:
        p = doc.add_paragraph("—")
        _apply_body(p, style)
    _blank(doc, 1)

    # === SECTIONS ===
    def add_h2(name: str) -> None:
        _heading_h2(doc, name)

    def add_body(text: str) -> None:
        _add_markdown_body(doc, text, style)
        _blank(doc, 1)

    add_h2("Introduction")
    add_body(str(summary.get("introduction") or ""))

    add_h2("Methods")
    add_body(str(summary.get("methods") or ""))

    # Results (subsections)
    add_h2("Results")
    results = summary.get("results") or []
    if isinstance(results, list) and results:
        for item in results:
            if not isinstance(item, dict):
                continue
            sec_title = str(item.get("section_title") or "").strip()
            sec_text = str(item.get("summary_text") or "").strip()

            if sec_title:
                h = doc.add_heading(sec_title, level=3)
                for r in h.runs:
                    r.font.size = Pt(12)
                h.paragraph_format.space_after = Pt(0)

            if sec_text:
                sec_text = _normalize_word_breaks(sec_text).strip()
                for chunk in sec_text.split("\n\n"):
                    p = doc.add_paragraph(chunk.strip())
                    _apply_body(p, style)
                _blank(doc, 1)
            else:
                p = doc.add_paragraph("—")
                _apply_body(p, style)
                _blank(doc, 1)
    else:
        add_body("")

    add_h2("Discussion")
    add_body(str(summary.get("discussion") or ""))

    add_h2("Figure narrative")
    add_body(str(summary.get("figure_narrative") or ""))

    # === ABBREVIATIONS ===
    _heading_h2(doc, "Abbreviations")
    abbr_text = str(summary.get("abbreviations") or "")
    abbr_pairs = _parse_abbreviation_pairs(abbr_text)
    if abbr_pairs:
        _abbrev_simple_table(doc, abbr_pairs)
    else:
        p = doc.add_paragraph("—")
        _apply_body(p, style)
    _blank(doc, 1)

    doc.save(str(docx_path))

